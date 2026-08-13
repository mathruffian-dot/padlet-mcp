"""產生 Phase 0 實測用的四種學生繳交樣本檔。

目的：測 get_board / get_attachment 對「照片 / 影片 / PDF / Word」四種
學生上傳附件分別回傳什麼。內容刻意寫成含錯誤的解題，讓後續視覺批改測試有東西可抓。
"""

import os
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
import docx
from docx.shared import Pt

OUT = Path(r"C:\2026Padlet_agent\phase0-samples")
OUT.mkdir(parents=True, exist_ok=True)

FONT_CANDIDATES = [
    (r"C:\Windows\Fonts\kaiu.ttf", None),       # 標楷體，純 TTF 最好處理
    (r"C:\Windows\Fonts\msjh.ttc", 0),          # 微軟正黑體
    (r"C:\Windows\Fonts\mingliu.ttc", 0),       # 細明體
    (r"C:\Windows\Fonts\simsun.ttc", 0),
]


def find_font():
    for path, idx in FONT_CANDIDATES:
        if os.path.exists(path):
            return path, idx
    return None, None


FONT_PATH, FONT_IDX = find_font()
if not FONT_PATH:
    sys.exit("找不到可用的中文字型，無法產生含中文的樣本檔")
print(f"使用字型：{FONT_PATH}")


def pil_font(size):
    kwargs = {"index": FONT_IDX} if FONT_IDX is not None else {}
    return ImageFont.truetype(FONT_PATH, size, **kwargs)


# ── 1. 照片：學生手寫解題（刻意寫錯，移項沒變號）──────────────────
def make_photo():
    W, H = 1200, 1600
    img = Image.new("RGB", (W, H), (252, 250, 244))          # 微黃紙感
    d = ImageDraw.Draw(img)
    for y in range(180, H, 90):                              # 筆記本橫線
        d.line([(80, y), (W - 80, y)], fill=(214, 224, 232), width=2)
    d.line([(150, 0), (150, H)], fill=(240, 180, 180), width=3)

    d.text((190, 70), "姓名：王小明   座號：07", font=pil_font(46), fill=(40, 40, 60))
    d.text((190, 195), "第 1 題　解 2x + 3 = 11", font=pil_font(58), fill=(20, 20, 40))

    steps = [
        "2x + 3 = 11",
        "2x = 11 + 3        ← 移項",
        "2x = 14",
        "x = 7",
    ]
    y = 340
    for s in steps:
        d.text((230, y), s, font=pil_font(54), fill=(25, 45, 120))
        y += 105
    d.text((230, y + 40), "答：x = 7", font=pil_font(58), fill=(25, 45, 120))
    path = OUT / "01-學生照片-手寫解題.png"
    img.save(path, "PNG")
    return path


# ── 2. 影片首格：之後用 ffmpeg 轉成 mp4 ─────────────────────────
def make_video_frame():
    W, H = 1280, 720
    img = Image.new("RGB", (W, H), (18, 24, 38))
    d = ImageDraw.Draw(img)
    d.text((90, 250), "學生口說解釋", font=pil_font(88), fill=(240, 244, 250))
    d.text((90, 380), "「我把 3 移到右邊，所以要加 3」", font=pil_font(46), fill=(255, 196, 120))
    d.text((90, 470), "（Phase 0 附件型態測試用樣本）", font=pil_font(34), fill=(140, 155, 180))
    path = OUT / "_video_frame.png"
    img.save(path, "PNG")
    return path


# ── 3. PDF：學生報告（兩頁）─────────────────────────────────────
def make_pdf():
    name = "SampleCJK"
    pdfmetrics.registerFont(
        TTFont(name, FONT_PATH, subfontIndex=FONT_IDX) if FONT_IDX is not None
        else TTFont(name, FONT_PATH)
    )
    path = OUT / "03-學生報告.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4

    c.setFont(name, 22)
    c.drawString(60, h - 80, "一元一次方程式　學習報告")
    c.setFont(name, 12)
    c.drawString(60, h - 110, "姓名：王小明　座號：07　日期：2026-08-13")
    body = [
        "",
        "一、我學到什麼",
        "　解方程式時要讓等號兩邊保持平衡，兩邊做一樣的事。",
        "",
        "二、我卡住的地方",
        "　移項的時候常常忘記變號，把 2x + 3 = 11 寫成 2x = 11 + 3。",
        "",
        "三、我的解法",
        "　2x + 3 = 11 → 2x = 11 + 3 → 2x = 14 → x = 7",
    ]
    y = h - 150
    for line in body:
        c.drawString(60, y, line)
        y -= 22
    c.showPage()

    c.setFont(name, 18)
    c.drawString(60, h - 80, "第二頁：練習題")
    c.setFont(name, 12)
    for i, q in enumerate(["3x - 5 = 10", "4(x + 2) = 20", "5x + 1 = 2x + 10"], 1):
        c.drawString(60, h - 120 - i * 26, f"{i}. {q}")
    c.save()
    return path


# ── 4. Word：學生作業 ───────────────────────────────────────────
def make_docx():
    doc = docx.Document()
    doc.add_heading("一元一次方程式　作業", level=1)
    p = doc.add_paragraph()
    p.add_run("姓名：王小明　座號：07").bold = True
    doc.add_paragraph("第 1 題　解 2x + 3 = 11")
    for line in ["2x + 3 = 11", "2x = 11 + 3   ← 這裡我不確定對不對", "2x = 14", "x = 7"]:
        doc.add_paragraph(line, style="List Number")
    doc.add_paragraph("我的問題：移項到底要不要變號？")
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(12)
    path = OUT / "04-學生作業.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for f in (make_photo(), make_video_frame(), make_pdf(), make_docx()):
        print(f"✅ {f.name}  ({f.stat().st_size:,} bytes)")
